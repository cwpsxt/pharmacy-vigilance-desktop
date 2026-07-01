"""
DeepSeek AI 助手模块 - 智能数据分析版
"""
import requests
import json
import re
import io
from datetime import datetime, timedelta
from flask import Blueprint, request, jsonify, send_file
from sqlalchemy import func, desc, case
from .auth import require_auth
from .db import db
from .models import AdverseReactionReport

ai_bp = Blueprint("ai", __name__)

DEEPSEEK_API_KEY = "sk-2cace93a7ac2415ea6a26c2c742a1eb6"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# 敏感字段列表（绝对禁止发送给AI）
SENSITIVE_FIELDS = ['report_code', 'medical_record_no', 'reporter_signature', 'import_batch_id']

SYSTEM_PROMPT = """你是一个专业的医院药学药物警戒助手。你可以帮助用户解答以下问题：
1. 药物不良反应相关知识
2. 药物警戒和安全用药指导
3. 数据分析方法和统计学知识
4. 系统使用帮助

注意：你不会访问或处理用户的任何业务数据，所有回答都基于通用医药知识。
请用简洁专业的中文回答用户问题。"""

# AI分析专用Prompt模板
PROMPTS = {
    'analyze': """你是医院药物警戒数据分析专家。请根据以下统计数据进行专业分析。

## 统计数据
{statistics}

请从以下角度分析并用Markdown格式输出：
1. **关键发现**：总结3-5个重要数据特征
2. **风险提示**：识别需要关注的异常点
3. **改进建议**：给出2-3条可操作建议

注意：数据已脱敏处理，不含任何患者隐私信息。请基于统计数据给出专业、简洁的分析。""",

    'trend': """你是医院药物警戒趋势分析专家。请根据以下历史数据进行趋势分析和预测。

## 历史数据
{trend_data}

请分析并用Markdown格式输出：
1. **趋势判断**：整体走势（上升/下降/平稳）
2. **预警信号**：是否存在需要警惕的趋势
3. **短期预测**：下一周期的预计情况
4. **应对建议**：具体措施建议

请用专业、简洁的语言回答。""",

    'report': """你是医院药物警戒报告撰写专家。请根据以下数据生成{report_type}。

## 报告数据
{report_data}

请生成结构清晰的Markdown格式报告，包含：
1. **概述**：总体情况简述
2. **数据分析**：关键指标解读
3. **重点关注**：需要注意的问题
4. **结论建议**：总结和改进建议

语言要求专业规范，适合作为正式报告使用。""",

    'advice': """你是临床药学专家。请根据以下药品不良反应统计数据，给出用药安全建议。

## 药品信息
药品名称：{drug_name}

## 本院统计数据
{statistics}

请提供专业的用药安全建议，包括：
1. **不良反应概述**：该药品常见不良反应说明
2. **高风险人群**：需要特别注意的患者群体
3. **预防措施**：降低不良反应风险的方法
4. **监测要点**：用药期间需要监测的指标

请用Markdown格式输出，语言专业简洁。"""
}


# ==================== 数据聚合函数（安全处理）====================

def validate_safe_data(data):
    """验证数据安全性，确保不含敏感字段"""
    data_str = json.dumps(data, ensure_ascii=False)
    for field in SENSITIVE_FIELDS:
        if field in data_str.lower():
            raise ValueError(f"数据包含敏感字段: {field}")
    return True


def get_safe_statistics(start_month=None, end_month=None):
    """获取安全的统计数据（不含敏感信息）"""
    query = AdverseReactionReport.query
    
    # 应用时间筛选
    if start_month:
        query = query.filter(
            func.strftime('%Y-%m', AdverseReactionReport.national_center_receive_time) >= start_month
        )
    if end_month:
        query = query.filter(
            func.strftime('%Y-%m', AdverseReactionReport.national_center_receive_time) <= end_month
        )
    
    # 1. 总量统计
    total = query.count()
    
    # 2. 类型分布
    type_stats = {
        '一般': query.filter(AdverseReactionReport.severity == '一般', 
                           AdverseReactionReport.report_type_new.is_(None)).count(),
        '严重': query.filter(AdverseReactionReport.severity == '严重',
                           AdverseReactionReport.report_type_new.is_(None)).count(),
        '新的+一般': query.filter(AdverseReactionReport.report_type_new == '新的',
                               AdverseReactionReport.severity == '一般').count(),
        '新的+严重': query.filter(AdverseReactionReport.report_type_new == '新的',
                               AdverseReactionReport.severity == '严重').count(),
    }
    
    # 3. 药品统计TOP10（仅药名+数量）
    drug_stats = db.session.query(
        AdverseReactionReport.generic_name,
        func.count(AdverseReactionReport.id).label('total'),
        func.sum(case((AdverseReactionReport.severity == '严重', 1), else_=0)).label('severe')
    ).group_by(AdverseReactionReport.generic_name)\
     .order_by(desc('total'))\
     .limit(10).all()
    
    drug_list = []
    for d in drug_stats:
        total_count = d.total or 0
        severe_count = d.severe or 0
        drug_list.append({
            'drug_name': d.generic_name,
            'count': total_count,
            'severe_count': severe_count,
            'severe_rate': f"{(severe_count/total_count*100):.1f}%" if total_count > 0 else "0%"
        })
    
    # 4. 不良反应类型统计TOP10
    reaction_stats = db.session.query(
        AdverseReactionReport.adverse_reaction_name,
        func.count(AdverseReactionReport.id).label('count')
    ).group_by(AdverseReactionReport.adverse_reaction_name)\
     .order_by(desc('count'))\
     .limit(10).all()
    
    reaction_list = [{'reaction_name': r.adverse_reaction_name, 'count': r.count} for r in reaction_stats]
    
    # 5. 科室/职业分布
    profession_stats = db.session.query(
        AdverseReactionReport.reporter_profession,
        func.count(AdverseReactionReport.id).label('count')
    ).group_by(AdverseReactionReport.reporter_profession)\
     .order_by(desc('count'))\
     .limit(10).all()
    
    profession_list = [{'profession': p.reporter_profession, 'count': p.count} for p in profession_stats]
    
    return {
        'period': f"{start_month or '最早'} 至 {end_month or '最新'}",
        'total': total,
        'type_distribution': type_stats,
        'top_drugs': drug_list,
        'top_reactions': reaction_list,
        'profession_distribution': profession_list
    }


def get_monthly_trend(months=6):
    """获取月度趋势数据"""
    end_date = datetime.now()
    start_date = end_date - timedelta(days=30 * months)
    
    monthly_data = db.session.query(
        func.strftime('%Y-%m', AdverseReactionReport.national_center_receive_time).label('month'),
        func.count(AdverseReactionReport.id).label('total'),
        func.sum(case((AdverseReactionReport.severity == '严重', 1), else_=0)).label('severe')
    ).filter(
        AdverseReactionReport.national_center_receive_time >= start_date
    ).group_by('month')\
     .order_by('month').all()
    
    trend_list = []
    for m in monthly_data:
        total = m.total or 0
        severe = m.severe or 0
        trend_list.append({
            'month': m.month,
            'total': total,
            'severe': severe,
            'severe_rate': f"{(severe/total*100):.1f}%" if total > 0 else "0%"
        })
    
    return trend_list


def get_drug_statistics(drug_name):
    """获取指定药品的统计数据"""
    query = AdverseReactionReport.query.filter(
        AdverseReactionReport.generic_name.like(f'%{drug_name}%')
    )
    
    total = query.count()
    severe = query.filter(AdverseReactionReport.severity == '严重').count()
    
    # 常见不良反应
    reactions = db.session.query(
        AdverseReactionReport.adverse_reaction_name,
        func.count(AdverseReactionReport.id).label('count')
    ).filter(
        AdverseReactionReport.generic_name.like(f'%{drug_name}%')
    ).group_by(AdverseReactionReport.adverse_reaction_name)\
     .order_by(desc('count'))\
     .limit(5).all()
    
    return {
        'drug_name': drug_name,
        'total_reactions': total,
        'severe_count': severe,
        'severe_rate': f"{(severe/total*100):.1f}%" if total > 0 else "0%",
        'common_reactions': [{'name': r.adverse_reaction_name, 'count': r.count} for r in reactions]
    }


def call_deepseek(prompt, temperature=0.7):
    """调用DeepSeek API"""
    response = requests.post(
        DEEPSEEK_API_URL,
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        },
        json={
            "model": "deepseek-v4-pro",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": temperature,
            "max_tokens": 3000,
            "stream": False
        },
        timeout=90
    )
    
    if response.status_code != 200:
        raise Exception(f"AI服务不可用 ({response.status_code})")
    
    result = response.json()
    return result.get("choices", [{}])[0].get("message", {}).get("content", "")


@ai_bp.route("/chat", methods=["POST"])
@require_auth
def chat():
    """AI 对话接口"""
    try:
        data = request.get_json(silent=True) or {}
        message = data.get("message", "").strip()
        history = data.get("history", [])
        
        if not message:
            return jsonify({"success": False, "message": "请输入问题"}), 400
        
        # 构建消息列表
        messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        
        # 添加历史消息（最多保留最近10轮）
        for msg in history[-20:]:
            messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
        
        # 添加当前消息
        messages.append({"role": "user", "content": message})
        
        # 调用 DeepSeek API
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-v4-pro",
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 2000,
                "stream": False
            },
            timeout=60
        )
        
        if response.status_code != 200:
            return jsonify({
                "success": False, 
                "message": f"AI 服务暂时不可用 ({response.status_code})"
            }), 500
        
        result = response.json()
        reply = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        
        if not reply:
            return jsonify({"success": False, "message": "AI 未能生成回复"}), 500
        
        return jsonify({
            "success": True,
            "reply": reply
        })
        
    except requests.exceptions.Timeout:
        return jsonify({"success": False, "message": "AI 响应超时，请稍后重试"}), 500
    except requests.exceptions.RequestException as e:
        return jsonify({"success": False, "message": "网络连接失败"}), 500
    except Exception as e:
        return jsonify({"success": False, "message": f"服务异常: {str(e)}"}), 500


# ==================== 智能分析接口 ====================

@ai_bp.route("/analyze", methods=["POST"])
@require_auth
def analyze_data():
    """数据洞察分析接口"""
    try:
        data = request.get_json(silent=True) or {}
        start_month = data.get("start_month")
        end_month = data.get("end_month")
        
        # 获取安全的统计数据
        statistics = get_safe_statistics(start_month, end_month)
        
        # 验证数据安全性
        validate_safe_data(statistics)
        
        # 格式化统计数据
        stats_text = f"""
时间范围：{statistics['period']}
报告总数：{statistics['total']}例

报告类型分布：
- 一般：{statistics['type_distribution']['一般']}例
- 严重：{statistics['type_distribution']['严重']}例
- 新的+一般：{statistics['type_distribution']['新的+一般']}例
- 新的+严重：{statistics['type_distribution']['新的+严重']}例

涉及药品TOP10：
{chr(10).join([f"- {d['drug_name']}：{d['count']}例（严重率{d['severe_rate']}）" for d in statistics['top_drugs'][:10]])}

常见不良反应TOP10：
{chr(10).join([f"- {r['reaction_name']}：{r['count']}例" for r in statistics['top_reactions'][:10]])}

报告人职业分布：
{chr(10).join([f"- {p['profession']}：{p['count']}例" for p in statistics['profession_distribution'][:5]])}
"""
        
        # 构建提示词并调用AI
        prompt = PROMPTS['analyze'].format(statistics=stats_text)
        reply = call_deepseek(prompt)
        
        return jsonify({
            "success": True,
            "reply": reply,
            "statistics": statistics
        })
        
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)}), 400
    except Exception as e:
        return jsonify({"success": False, "message": f"分析失败: {str(e)}"}), 500


@ai_bp.route("/trend", methods=["POST"])
@require_auth
def analyze_trend():
    """趋势预测预警接口"""
    try:
        data = request.get_json(silent=True) or {}
        months = data.get("months", 6)
        
        # 获取趋势数据
        trend_data = get_monthly_trend(months)
        
        if not trend_data:
            return jsonify({"success": False, "message": "暂无足够的历史数据进行趋势分析"}), 400
        
        # 验证数据安全性
        validate_safe_data(trend_data)
        
        # 格式化趋势数据
        trend_text = f"""
分析周期：近{months}个月

月度数据：
{chr(10).join([f"- {t['month']}：共{t['total']}例，严重{t['severe']}例（{t['severe_rate']}）" for t in trend_data])}

趋势概要：
- 起始月报告量：{trend_data[0]['total'] if trend_data else 0}例
- 最近月报告量：{trend_data[-1]['total'] if trend_data else 0}例
- 变化幅度：{((trend_data[-1]['total'] - trend_data[0]['total']) / trend_data[0]['total'] * 100):.1f}% （较首月）
"""
        
        # 调用AI分析
        prompt = PROMPTS['trend'].format(trend_data=trend_text)
        reply = call_deepseek(prompt)
        
        return jsonify({
            "success": True,
            "reply": reply,
            "trend_data": trend_data
        })
        
    except Exception as e:
        return jsonify({"success": False, "message": f"趋势分析失败: {str(e)}"}), 500


@ai_bp.route("/report", methods=["POST"])
@require_auth
def generate_report():
    """智能报告生成接口"""
    try:
        data = request.get_json(silent=True) or {}
        report_type = data.get("report_type", "月度分析报告")
        start_month = data.get("start_month")
        end_month = data.get("end_month")
        
        # 获取统计数据
        statistics = get_safe_statistics(start_month, end_month)
        trend_data = get_monthly_trend(6)
        
        # 验证数据安全性
        validate_safe_data(statistics)
        
        # 格式化报告数据
        report_data = f"""
报告类型：{report_type}
统计周期：{statistics['period']}

一、基础数据
- 报告总数：{statistics['total']}例
- 一般反应：{statistics['type_distribution']['一般']}例
- 严重反应：{statistics['type_distribution']['严重']}例
- 新的一般：{statistics['type_distribution']['新的+一般']}例
- 新的严重：{statistics['type_distribution']['新的+严重']}例

二、药品分析（TOP10）
{chr(10).join([f"{i+1}. {d['drug_name']}：{d['count']}例，严重率{d['severe_rate']}" for i, d in enumerate(statistics['top_drugs'][:10])])}

三、不良反应类型分析（TOP10）
{chr(10).join([f"{i+1}. {r['reaction_name']}：{r['count']}例" for i, r in enumerate(statistics['top_reactions'][:10])])}

四、月度趋势
{chr(10).join([f"- {t['month']}：{t['total']}例" for t in trend_data[-6:]])}
"""
        
        # 调用AI生成报告
        prompt = PROMPTS['report'].format(report_type=report_type, report_data=report_data)
        reply = call_deepseek(prompt, temperature=0.5)
        
        return jsonify({
            "success": True,
            "reply": reply,
            "statistics": statistics
        })
        
    except Exception as e:
        return jsonify({"success": False, "message": f"报告生成失败: {str(e)}"}), 500


@ai_bp.route("/advice", methods=["POST"])
@require_auth
def drug_advice():
    """用药安全建议接口"""
    try:
        data = request.get_json(silent=True) or {}
        drug_name = data.get("drug_name", "").strip()
        
        if not drug_name:
            return jsonify({"success": False, "message": "请提供药品名称"}), 400
        
        # 获取药品统计数据
        drug_stats = get_drug_statistics(drug_name)
        
        if drug_stats['total_reactions'] == 0:
            return jsonify({
                "success": True,
                "reply": f"系统中暂无「{drug_name}」相关的不良反应记录。\n\n如需了解该药品的一般用药安全信息，请直接向我提问。",
                "statistics": drug_stats
            })
        
        # 验证数据安全性
        validate_safe_data(drug_stats)
        
        # 格式化统计数据
        stats_text = f"""
不良反应报告总数：{drug_stats['total_reactions']}例
严重反应数：{drug_stats['severe_count']}例
严重反应率：{drug_stats['severe_rate']}

常见不良反应：
{chr(10).join([f"- {r['name']}：{r['count']}例" for r in drug_stats['common_reactions']])}
"""
        
        # 调用AI生成建议
        prompt = PROMPTS['advice'].format(drug_name=drug_name, statistics=stats_text)
        reply = call_deepseek(prompt)
        
        return jsonify({
            "success": True,
            "reply": reply,
            "statistics": drug_stats
        })
        
    except Exception as e:
        return jsonify({"success": False, "message": f"获取建议失败: {str(e)}"}), 500


@ai_bp.route("/help", methods=["GET"])
@require_auth
def ai_help():
    """获取AI助手帮助信息"""
    help_text = """## 🤖 AI药物警戒助手使用指南

### 快捷命令
| 命令 | 功能 | 示例 |
|-----|------|-----|
| `/分析` | 数据洞察分析 | `/分析` 或 `/分析 2024-01 2024-06` |
| `/趋势` | 趋势预测预警 | `/趋势` 或 `/趋势 12` (近12个月) |
| `/报告` | 生成分析报告 | `/报告 月度` 或 `/报告 季度` |
| `/建议 药品名` | 用药安全建议 | `/建议 阿莫西林` |
| `/帮助` | 显示此帮助 | `/帮助` |

### 功能说明
1. **数据洞察**：分析不良反应数据特征，发现规律和异常
2. **趋势预警**：预测报告趋势，识别预警信号
3. **智能报告**：自动生成专业的分析报告
4. **用药建议**：基于本院数据提供用药安全建议

### 数据安全
- ✅ 所有分析基于统计聚合数据
- ✅ 不会发送任何患者隐私信息
- ✅ 病历号、报告编号、报告人等信息严格保护

您也可以直接输入问题，我会基于专业知识为您解答。"""
    
    return jsonify({
        "success": True,
        "help": help_text
    })


@ai_bp.route("/export-word", methods=["POST"])
@require_auth
def export_report_word():
    """导出报告为Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        
        data = request.get_json(silent=True) or {}
        content = data.get("content", "")
        title = data.get("title", "药品不良反应分析报告")
        
        if not content:
            return jsonify({"success": False, "message": "报告内容为空"}), 400
        
        # 清理Markdown符号的辅助函数
        def clean_markdown(text):
            # 移除标题符号
            text = re.sub(r'^#{1,6}\s*', '', text)
            # 移除粗体符号但保留内容
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # 移除斜体符号
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            # 移除代码符号
            text = re.sub(r'`(.+?)`', r'\1', text)
            # 移除emoji前缀（如📊）
            text = re.sub(r'^[📊📈📝💊❓⚠️✅🔍📋🏥💉🩺🔹▪️•]+\s*', '', text)
            return text.strip()
        
        # 提取粗体文本位置
        def parse_bold_text(text):
            """返回 [(text, is_bold), ...]"""
            parts = []
            last_end = 0
            for match in re.finditer(r'\*\*(.+?)\*\*', text):
                if match.start() > last_end:
                    parts.append((text[last_end:match.start()], False))
                parts.append((match.group(1), True))
                last_end = match.end()
            if last_end < len(text):
                parts.append((text[last_end:], False))
            return parts if parts else [(text, False)]
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.styles['Normal'].font.size = Pt(11)
        
        # 设置文档标题
        title_para = doc.add_heading('', 0)
        title_run = title_para.add_run(title)
        title_run.font.name = '微软雅黑'
        title_run.font.size = Pt(22)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _now = datetime.now()
        time_run = time_para.add_run(f"生成时间：{_now.year}年{_now.month:02d}月{_now.day:02d}日 {_now.hour:02d}:{_now.minute:02d}")
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加空行分隔
        doc.add_paragraph()
        
        # 解析Markdown内容并添加到文档
        lines = content.split('\n')
        current_list = []
        in_table = False
        table_rows = []
        
        for i, line in enumerate(lines):
            original_line = line
            line = line.strip()
            
            if not line:
                # 处理累积的列表
                if current_list:
                    for item in current_list:
                        clean_item = clean_markdown(item)
                        p = doc.add_paragraph(style='List Bullet')
                        # 处理粗体
                        for text, is_bold in parse_bold_text(item):
                            clean_text = clean_markdown(text) if not is_bold else text
                            run = p.add_run(clean_text)
                            run.bold = is_bold
                            run.font.name = '微软雅黑'
                    current_list = []
                # 处理表格
                if in_table and table_rows:
                    _add_table_to_doc(doc, table_rows)
                    table_rows = []
                    in_table = False
                continue
            
            # 处理表格
            if line.startswith('|') and line.endswith('|'):
                if '---' in line or ':-' in line or '-:' in line:
                    continue  # 跳过分隔行
                in_table = True
                cells = [clean_markdown(c.strip()) for c in line.split('|')[1:-1]]
                table_rows.append(cells)
                continue
            
            # 如果之前在表格中，现在不是表格行了，先输出表格
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            
            # 处理标题（支持1-6级）
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                if current_list:
                    for item in current_list:
                        p = doc.add_paragraph(clean_markdown(item), style='List Bullet')
                    current_list = []
                level = len(heading_match.group(1))
                heading_text = clean_markdown(heading_match.group(2))
                # Word只支持1-9级标题，我们用1-3级
                word_level = min(level, 3)
                h = doc.add_heading('', level=word_level)
                run = h.add_run(heading_text)
                run.font.name = '微软雅黑'
                if level == 1:
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 51, 102)
                elif level == 2:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 82, 155)
                else:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(51, 51, 51)
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                current_list.append(line[2:].strip())
            elif re.match(r'^\d+\. ', line):
                current_list.append(re.sub(r'^\d+\. ', '', line).strip())
            # 普通段落
            else:
                if current_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        for text, is_bold in parse_bold_text(item):
                            clean_text = clean_markdown(text) if not is_bold else text
                            run = p.add_run(clean_text)
                            run.bold = is_bold
                            run.font.name = '微软雅黑'
                    current_list = []
                
                # 添加段落，处理粗体
                p = doc.add_paragraph()
                for text, is_bold in parse_bold_text(line):
                    clean_text = clean_markdown(text) if not is_bold else text
                    run = p.add_run(clean_text)
                    run.bold = is_bold
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(11)
        
        # 处理剩余列表
        if current_list:
            for item in current_list:
                p = doc.add_paragraph(style='List Bullet')
                for text, is_bold in parse_bold_text(item):
                    clean_text = clean_markdown(text) if not is_bold else text
                    run = p.add_run(clean_text)
                    run.bold = is_bold
                    run.font.name = '微软雅黑'
        
        # 处理剩余表格
        if table_rows:
            _add_table_to_doc(doc, table_rows)
        
        # 添加页脚说明
        doc.add_paragraph()
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run("本报告由AI药物警戒助手自动生成，数据已脱敏处理，仅供参考。")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.font.name = '微软雅黑'
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存到内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # 生成文件名
        filename = f"药品不良反应分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except ImportError:
        return jsonify({"success": False, "message": "服务器未安装python-docx库"}), 500
    except Exception as e:
        return jsonify({"success": False, "message": f"导出失败: {str(e)}"}), 500


def _add_table_to_doc(doc, rows):
    """添加表格到文档"""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                # 设置字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run.font.size = Pt(10)
                # 首行加粗（表头）
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()  # 表格后空行


@ai_bp.route("/chart", methods=["POST"])
@require_auth
def generate_chart():
    """根据用户请求生成图表"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from matplotlib import font_manager
        import base64
        
        # 设置中文字体
        plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'SimHei', 'Microsoft YaHei']
        plt.rcParams['axes.unicode_minus'] = False
        
        data = request.get_json(silent=True) or {}
        chart_type = data.get("chart_type", "bar")  # bar, pie, line
        dataset = data.get("dataset", "report")  # report, workload
        start_month = data.get("start_month")
        end_month = data.get("end_month")
        dimension = data.get("dimension", "report_type")  # report_type, severity, drug, reaction
        title = data.get("title", "统计图表")
        
        # 构建查询
        query = AdverseReactionReport.query
        
        if start_month:
            query = query.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
        if end_month:
            query = query.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
        
        # 根据维度统计数据
        chart_data = {}
        
        if dimension == "report_type":
            # 报告类型统计
            types = ['新的', '严重', '一般', '新的+严重', '新的+一般']
            for t in types:
                if '+' in t:
                    parts = t.split('+')
                    count = query.filter(
                        AdverseReactionReport.report_type_new == parts[0],
                        AdverseReactionReport.severity == parts[1]
                    ).count()
                elif t == '新的':
                    count = query.filter(AdverseReactionReport.report_type_new == '新的').count()
                else:
                    count = query.filter(AdverseReactionReport.severity == t).count()
                if count > 0:
                    chart_data[t] = count
                    
        elif dimension == "severity":
            # 严重程度统计
            for sev in ['一般', '严重']:
                count = query.filter(AdverseReactionReport.severity == sev).count()
                if count > 0:
                    chart_data[sev] = count
                    
        elif dimension == "drug":
            # 药品TOP10统计
            drug_stats = db.session.query(
                AdverseReactionReport.generic_name,
                func.count(AdverseReactionReport.id).label('count')
            )
            if start_month:
                drug_stats = drug_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
            if end_month:
                drug_stats = drug_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
            drug_stats = drug_stats.group_by(AdverseReactionReport.generic_name)\
                .order_by(desc('count')).limit(10).all()
            for d in drug_stats:
                if d.generic_name:
                    chart_data[d.generic_name[:10]] = d.count
                    
        elif dimension == "reaction":
            # 不良反应类型TOP10
            reaction_stats = db.session.query(
                AdverseReactionReport.adverse_reaction_name,
                func.count(AdverseReactionReport.id).label('count')
            ).filter(
                AdverseReactionReport.adverse_reaction_name != None,
                AdverseReactionReport.adverse_reaction_name != ''
            )
            if start_month:
                reaction_stats = reaction_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
            if end_month:
                reaction_stats = reaction_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
            reaction_stats = reaction_stats.group_by(AdverseReactionReport.adverse_reaction_name)\
                .order_by(desc('count')).limit(10).all()
            for r in reaction_stats:
                if r.adverse_reaction_name:
                    chart_data[r.adverse_reaction_name[:10]] = r.count
                    
        elif dimension == "monthly":
            # 月度趋势
            monthly_stats = db.session.query(
                func.strftime('%Y-%m', AdverseReactionReport.national_center_receive_time).label('month'),
                func.count(AdverseReactionReport.id).label('count')
            )
            if start_month:
                monthly_stats = monthly_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
            if end_month:
                monthly_stats = monthly_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
            monthly_stats = monthly_stats.group_by('month').order_by('month').all()
            for m in monthly_stats:
                if m.month:
                    chart_data[m.month] = m.count
        
        if not chart_data:
            return jsonify({"success": False, "message": "所选时间范围内没有数据"}), 400
        
        # 生成图表
        fig, ax = plt.subplots(figsize=(10, 6))
        labels = list(chart_data.keys())
        values = list(chart_data.values())
        colors = ['#1890ff', '#52c41a', '#faad14', '#f5222d', '#722ed1', 
                  '#13c2c2', '#eb2f96', '#fa8c16', '#a0d911', '#2f54eb']
        
        if chart_type == "bar":
            bars = ax.bar(labels, values, color=colors[:len(labels)])
            ax.set_ylabel('数量', fontsize=12)
            # 在柱状图上显示数值
            for bar, val in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                       str(val), ha='center', va='bottom', fontsize=10)
            plt.xticks(rotation=45, ha='right')
            
        elif chart_type == "pie":
            ax.pie(values, labels=labels, autopct='%1.1f%%', colors=colors[:len(labels)],
                   startangle=90, explode=[0.02]*len(labels))
            ax.axis('equal')
            
        elif chart_type == "line":
            ax.plot(labels, values, marker='o', linewidth=2, markersize=8, color='#1890ff')
            ax.fill_between(labels, values, alpha=0.3, color='#1890ff')
            ax.set_ylabel('数量', fontsize=12)
            for i, val in enumerate(values):
                ax.annotate(str(val), (labels[i], val), textcoords="offset points", 
                           xytext=(0,10), ha='center', fontsize=10)
            plt.xticks(rotation=45, ha='right')
        
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        plt.tight_layout()
        
        # 保存为base64图片
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
        plt.close(fig)
        
        # 生成文字描述
        total = sum(values)
        max_item = max(chart_data.items(), key=lambda x: x[1])
        description = f"共统计{total}条记录，其中「{max_item[0]}」最多，有{max_item[1]}条，占比{max_item[1]/total*100:.1f}%。"
        
        return jsonify({
            "success": True,
            "chart_image": f"data:image/png;base64,{img_base64}",
            "description": description,
            "data": chart_data
        })
        
    except Exception as e:
        return jsonify({"success": False, "message": f"生成图表失败: {str(e)}"}), 500


@ai_bp.route("/smart-query", methods=["POST"])
@require_auth  
def smart_query():
    """智能查询 - 解析用户自然语言请求并返回统计数据用于ECharts渲染"""
    try:
        data = request.get_json(silent=True) or {}
        query_text = data.get("query", "").strip()
        dataset = data.get("dataset", "report")
        start_month = data.get("start_month")
        end_month = data.get("end_month")
        
        if not query_text:
            return jsonify({"success": False, "message": "请输入查询内容"}), 400
        
        # 解析用户意图
        chart_type = "bar"
        dimension = "report_type"
        title = "统计分析"
        
        query_lower = query_text.lower()
        
        # 解析图表类型
        if "饼" in query_text or "pie" in query_lower or "占比" in query_text or "比例" in query_text:
            chart_type = "pie"
        elif "折线" in query_text or "趋势" in query_text or "line" in query_lower or "变化" in query_text:
            chart_type = "line"
        elif "柱" in query_text or "bar" in query_lower:
            chart_type = "bar"
        
        # 解析统计维度
        if "报告类型" in query_text or "类型统计" in query_text or "类型分布" in query_text:
            dimension = "report_type"
            title = "报告类型统计"
        elif "严重" in query_text and ("程度" in query_text or "级别" in query_text or "分布" in query_text):
            dimension = "severity"
            title = "严重程度分布"
        elif "药品" in query_text or "药物" in query_text or "用药" in query_text:
            dimension = "drug"
            title = "药品TOP10统计"
        elif "反应" in query_text or "症状" in query_text or "表现" in query_text:
            dimension = "reaction"
            title = "不良反应类型TOP10"
        elif "月" in query_text and ("趋势" in query_text or "变化" in query_text or "统计" in query_text):
            dimension = "monthly"
            title = "月度趋势统计"
            chart_type = "line"
        
        # 从查询中提取时间（覆盖传入参数）
        year_month_pattern = r'(\d{4})年(\d{1,2})月'
        matches = re.findall(year_month_pattern, query_text)
        if len(matches) >= 2:
            start_month = f"{matches[0][0]}-{int(matches[0][1]):02d}"
            end_month = f"{matches[1][0]}-{int(matches[1][1]):02d}"
        elif len(matches) == 1:
            start_month = f"{matches[0][0]}-{int(matches[0][1]):02d}"
            end_month = start_month
        
        # 构建标题
        time_str = ""
        if start_month and end_month:
            if start_month == end_month:
                time_str = f"（{start_month}）"
            else:
                time_str = f"（{start_month} 至 {end_month}）"
        title = f"{title}{time_str}"
        
        # 获取统计数据
        chart_data = get_chart_data(dimension, start_month, end_month)
        
        if not chart_data:
            return jsonify({"success": False, "message": "所选时间范围内没有数据"}), 400
        
        # 生成描述
        total = sum(chart_data.values())
        max_item = max(chart_data.items(), key=lambda x: x[1])
        description = f"共统计{total}条记录，其中「{max_item[0]}」最多，有{max_item[1]}条，占比{max_item[1]/total*100:.1f}%。"
        
        return jsonify({
            "success": True,
            "chart_type": chart_type,
            "title": title,
            "data": chart_data,
            "labels": list(chart_data.keys()),
            "values": list(chart_data.values()),
            "description": description,
            "analysis": f"根据您的查询「{query_text}」，已生成{title}。{description}"
        })
        
    except Exception as e:
        return jsonify({"success": False, "message": f"智能查询失败: {str(e)}"}), 500


def get_chart_data(dimension, start_month, end_month):
    """获取图表统计数据（过滤并用数据，仅保留怀疑数据，按报告编码去重）"""
    print(f"[DEBUG] get_chart_data: dimension={dimension}, start_month={start_month}, end_month={end_month}")
    
    chart_data = {}
    
    # 构建基础过滤条件：只保留"怀疑"数据，过滤掉"并用"数据
    base_filter = AdverseReactionReport.suspect_concurrent == '怀疑'
    
    # 构建时间过滤条件
    time_filters = [base_filter]
    if start_month:
        time_filters.append(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
    if end_month:
        time_filters.append(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
    
    if dimension == "report_type":
        # 报告类型统计（按report_code去重）
        # 统计新的报告数
        new_query = db.session.query(func.count(func.distinct(AdverseReactionReport.report_code)))\
            .filter(AdverseReactionReport.report_type_new == '新的')
        for f in time_filters:
            new_query = new_query.filter(f)
        new_count = new_query.scalar() or 0
        
        # 统计一般报告数
        general_query = db.session.query(func.count(func.distinct(AdverseReactionReport.report_code)))\
            .filter(AdverseReactionReport.severity == '一般')
        for f in time_filters:
            general_query = general_query.filter(f)
        general_count = general_query.scalar() or 0
        
        # 统计严重报告数
        severe_query = db.session.query(func.count(func.distinct(AdverseReactionReport.report_code)))\
            .filter(AdverseReactionReport.severity == '严重')
        for f in time_filters:
            severe_query = severe_query.filter(f)
        severe_count = severe_query.scalar() or 0
        
        if new_count > 0:
            chart_data['新的'] = new_count
        if general_count > 0:
            chart_data['一般'] = general_count
        if severe_count > 0:
            chart_data['严重'] = severe_count
        
        print(f"[DEBUG] 报告类型统计(去重): 新的={new_count}, 一般={general_count}, 严重={severe_count}")
                
    elif dimension == "severity":
        for sev in ['一般', '严重']:
            query = db.session.query(func.count(func.distinct(AdverseReactionReport.report_code)))\
                .filter(AdverseReactionReport.severity == sev)
            for f in time_filters:
                query = query.filter(f)
            count = query.scalar() or 0
            if count > 0:
                chart_data[sev] = count
                
    elif dimension == "drug":
        # 药品统计：过滤并用，按report_code去重
        drug_query = db.session.query(
            AdverseReactionReport.generic_name,
            func.count(func.distinct(AdverseReactionReport.report_code)).label('count')
        ).filter(
            AdverseReactionReport.suspect_concurrent == '怀疑',
            AdverseReactionReport.generic_name != None,
            AdverseReactionReport.generic_name != ''
        )
        # 应用时间过滤
        if start_month:
            drug_query = drug_query.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
        if end_month:
            drug_query = drug_query.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
        
        drug_stats = drug_query.group_by(AdverseReactionReport.generic_name)\
            .order_by(desc('count')).limit(10).all()
        
        print(f"[DEBUG] drug_stats结果数量: {len(drug_stats)}")
        for d in drug_stats:
            print(f"[DEBUG] 药品: {d.generic_name}, 数量: {d.count}")
            if d.generic_name and d.generic_name.strip():
                chart_data[d.generic_name[:10]] = d.count
                
    elif dimension == "reaction":
        # 不良反应统计：过滤并用，按report_code去重
        reaction_stats = db.session.query(
            AdverseReactionReport.adverse_reaction_name,
            func.count(func.distinct(AdverseReactionReport.report_code)).label('count')
        ).filter(
            AdverseReactionReport.suspect_concurrent == '怀疑',
            AdverseReactionReport.adverse_reaction_name != None,
            AdverseReactionReport.adverse_reaction_name != ''
        )
        if start_month:
            reaction_stats = reaction_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
        if end_month:
            reaction_stats = reaction_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
        reaction_stats = reaction_stats.group_by(AdverseReactionReport.adverse_reaction_name)\
            .order_by(desc('count')).limit(10).all()
        for r in reaction_stats:
            if r.adverse_reaction_name and r.adverse_reaction_name.strip():
                chart_data[r.adverse_reaction_name[:10]] = r.count
                
    elif dimension == "monthly":
        # 月度统计：过滤并用，按report_code去重
        monthly_stats = db.session.query(
            func.strftime('%Y-%m', AdverseReactionReport.national_center_receive_time).label('month'),
            func.count(func.distinct(AdverseReactionReport.report_code)).label('count')
        ).filter(AdverseReactionReport.suspect_concurrent == '怀疑')
        if start_month:
            monthly_stats = monthly_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
        if end_month:
            monthly_stats = monthly_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
        monthly_stats = monthly_stats.group_by('month').order_by('month').all()
        for m in monthly_stats:
            if m.month:
                chart_data[m.month] = m.count
    
    print(f"[DEBUG] 最终chart_data: {chart_data}")
    return chart_data


def generate_chart_internal(chart_type, dataset, start_month, end_month, dimension, title):
    """内部图表生成函数"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import base64
        
        plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'SimHei', 'Microsoft YaHei']
        plt.rcParams['axes.unicode_minus'] = False
        
        # 构建查询
        query = AdverseReactionReport.query
        
        if start_month:
            query = query.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
        if end_month:
            query = query.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
        
        chart_data = {}
        
        if dimension == "report_type":
            types = ['新的', '严重', '一般']
            for t in types:
                if t == '新的':
                    count = query.filter(AdverseReactionReport.report_type_new == '新的').count()
                else:
                    count = query.filter(AdverseReactionReport.severity == t).count()
                if count > 0:
                    chart_data[t] = count
                    
        elif dimension == "severity":
            for sev in ['一般', '严重']:
                count = query.filter(AdverseReactionReport.severity == sev).count()
                if count > 0:
                    chart_data[sev] = count
                    
        elif dimension == "drug":
            drug_stats = db.session.query(
                AdverseReactionReport.generic_name,
                func.count(AdverseReactionReport.id).label('count')
            )
            if start_month:
                drug_stats = drug_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
            if end_month:
                drug_stats = drug_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
            drug_stats = drug_stats.group_by(AdverseReactionReport.generic_name)\
                .order_by(desc('count')).limit(10).all()
            for d in drug_stats:
                if d.generic_name:
                    chart_data[d.generic_name[:8]] = d.count
                    
        elif dimension == "reaction":
            reaction_stats = db.session.query(
                AdverseReactionReport.adverse_reaction_name,
                func.count(AdverseReactionReport.id).label('count')
            ).filter(
                AdverseReactionReport.adverse_reaction_name != None,
                AdverseReactionReport.adverse_reaction_name != ''
            )
            if start_month:
                reaction_stats = reaction_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
            if end_month:
                reaction_stats = reaction_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
            reaction_stats = reaction_stats.group_by(AdverseReactionReport.adverse_reaction_name)\
                .order_by(desc('count')).limit(10).all()
            for r in reaction_stats:
                if r.adverse_reaction_name:
                    chart_data[r.adverse_reaction_name[:8]] = r.count
                    
        elif dimension == "monthly":
            monthly_stats = db.session.query(
                func.strftime('%Y-%m', AdverseReactionReport.national_center_receive_time).label('month'),
                func.count(AdverseReactionReport.id).label('count')
            )
            if start_month:
                monthly_stats = monthly_stats.filter(AdverseReactionReport.national_center_receive_time >= f"{start_month}-01")
            if end_month:
                monthly_stats = monthly_stats.filter(AdverseReactionReport.national_center_receive_time <= f"{end_month}-31")
            monthly_stats = monthly_stats.group_by('month').order_by('month').all()
            for m in monthly_stats:
                if m.month:
                    chart_data[m.month] = m.count
        
        if not chart_data:
            return {"success": False, "message": "所选时间范围内没有数据"}
        
        # 生成图表
        fig, ax = plt.subplots(figsize=(10, 6))
        labels = list(chart_data.keys())
        values = list(chart_data.values())
        colors = ['#1890ff', '#52c41a', '#faad14', '#f5222d', '#722ed1', 
                  '#13c2c2', '#eb2f96', '#fa8c16', '#a0d911', '#2f54eb']
        
        if chart_type == "bar":
            bars = ax.bar(labels, values, color=colors[:len(labels)])
            ax.set_ylabel('数量', fontsize=12)
            for bar, val in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                       str(val), ha='center', va='bottom', fontsize=10)
            plt.xticks(rotation=45, ha='right')
            
        elif chart_type == "pie":
            ax.pie(values, labels=labels, autopct='%1.1f%%', colors=colors[:len(labels)],
                   startangle=90)
            ax.axis('equal')
            
        elif chart_type == "line":
            ax.plot(labels, values, marker='o', linewidth=2, markersize=8, color='#1890ff')
            ax.fill_between(labels, values, alpha=0.3, color='#1890ff')
            ax.set_ylabel('数量', fontsize=12)
            for i, val in enumerate(values):
                ax.annotate(str(val), (labels[i], val), textcoords="offset points", 
                           xytext=(0,10), ha='center', fontsize=10)
            plt.xticks(rotation=45, ha='right')
        
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        plt.tight_layout()
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
        plt.close(fig)
        
        total = sum(values)
        max_item = max(chart_data.items(), key=lambda x: x[1])
        description = f"共统计{total}条记录，其中「{max_item[0]}」最多，有{max_item[1]}条，占比{max_item[1]/total*100:.1f}%。"
        
        return {
            "success": True,
            "chart_image": f"data:image/png;base64,{img_base64}",
            "description": description,
            "data": chart_data
        }
        
    except Exception as e:
        return {"success": False, "message": str(e)}
