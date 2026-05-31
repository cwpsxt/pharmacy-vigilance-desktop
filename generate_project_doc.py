#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
药品不良反应数据分析系统 - 项目技术文档生成脚本
生成Word格式的系统说明文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_document():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 设置标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Microsoft YaHei'
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    
    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(12)
    
    return doc

def add_cover_page(doc):
    """添加封面"""
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('医院药学药物警戒与运营绩效\n可视化分析系统')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('系统技术架构与功能说明文档')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    for _ in range(8):
        doc.add_paragraph()
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'{datetime.datetime.now().strftime("%Y年%m月")}')
    run.font.size = Pt(14)
    
    doc.add_page_break()

def add_toc(doc):
    """添加目录"""
    doc.add_heading('目  录', level=1)
    
    toc_items = [
        ('一、系统概述', 1),
        ('二、系统架构设计', 1),
        ('    2.1 整体架构', 2),
        ('    2.2 技术栈选型', 2),
        ('    2.3 跨平台与离线运行', 2),
        ('三、核心功能模块', 1),
        ('    3.1 数据管理模块', 2),
        ('    3.2 可视化图表分析', 2),
        ('    3.3 AI智能助手', 2),
        ('    3.4 工作量统计分析', 2),
        ('四、创新功能亮点', 1),
        ('    4.1 AI驱动的智能分析', 2),
        ('    4.2 数据隐私保护机制', 2),
        ('    4.3 交互式可视化图表', 2),
        ('五、数据安全与隐私保护', 1),
        ('六、系统部署与运行', 1),
        ('七、总结与展望', 1),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        if level == 1:
            p.paragraph_format.left_indent = Cm(0)
        else:
            p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_page_break()

def add_section_overview(doc):
    """第一章：系统概述"""
    doc.add_heading('一、系统概述', level=1)
    
    doc.add_paragraph(
        '医院药学药物警戒与运营绩效可视化分析系统是一款专为医院药学部门设计的智能化数据分析平台。'
        '系统基于现代化的Electron桌面应用框架与Python后端服务构建，实现了药品不良反应数据的'
        '智能化采集、分析、可视化展示及报告生成等核心功能。'
    )
    
    doc.add_heading('1.1 系统定位', level=2)
    doc.add_paragraph(
        '本系统定位为医院药物警戒工作的核心支撑平台，旨在：\n'
        '• 提升药品不良反应监测工作效率\n'
        '• 实现数据驱动的科学决策支持\n'
        '• 保障患者用药安全\n'
        '• 满足国家药品不良反应监测法规要求'
    )
    
    doc.add_heading('1.2 适用场景', level=2)
    doc.add_paragraph(
        '• 医院药学部日常药物警戒工作\n'
        '• 药品不良反应月度/季度/年度报告编制\n'
        '• 药品安全风险预警与趋势分析\n'
        '• 科室绩效考核与工作量统计\n'
        '• 药品分类监测与统计分析'
    )

def add_section_architecture(doc):
    """第二章：系统架构设计"""
    doc.add_heading('二、系统架构设计', level=1)
    
    doc.add_heading('2.1 整体架构', level=2)
    doc.add_paragraph(
        '系统采用前后端分离的架构设计，前端基于Electron实现跨平台桌面应用，'
        '后端采用Python Flask框架提供RESTful API服务。整体架构如下：'
    )
    
    # 架构图描述
    arch_para = doc.add_paragraph()
    arch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = arch_para.add_run(
        '┌─────────────────────────────────────────────────────────┐\n'
        '│                    用户界面层 (Electron + Vue.js)              │\n'
        '├─────────────────────────────────────────────────────────┤\n'
        '│   登录认证  │  数据展示  │  图表可视化  │  AI对话交互   │\n'
        '├─────────────────────────────────────────────────────────┤\n'
        '│                    业务逻辑层 (Python Flask)                   │\n'
        '├─────────────────────────────────────────────────────────┤\n'
        '│  数据管理  │  统计分析  │  报告生成  │  AI智能服务   │\n'
        '├─────────────────────────────────────────────────────────┤\n'
        '│                    数据存储层 (SQLite)                          │\n'
        '└─────────────────────────────────────────────────────────┘'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('2.2 技术栈选型', level=2)
    
    # 技术栈表格
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['层级', '技术组件', '说明']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E0E7FF')
    
    data = [
        ('前端框架', 'Electron 28.x', '跨平台桌面应用框架'),
        ('UI框架', 'Vue.js 3 + Bootstrap 5', '响应式用户界面'),
        ('图表库', 'ECharts 5.x', '交互式数据可视化'),
        ('后端框架', 'Python Flask', '轻量级Web服务框架'),
        ('数据库', 'SQLite', '嵌入式关系数据库'),
        ('AI服务', 'DeepSeek API', '大语言模型智能分析'),
        ('文档生成', 'python-docx', 'Word文档自动生成'),
    ]
    
    for i, (col1, col2, col3) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = col1
        row.cells[1].text = col2
        row.cells[2].text = col3
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 跨平台与离线运行', level=2)
    
    doc.add_heading('跨平台支持', level=3)
    doc.add_paragraph(
        '系统基于Electron框架开发，原生支持以下操作系统：\n'
        '• Windows 10/11 (64位)\n'
        '• macOS 10.15+ (Intel/Apple Silicon)\n'
        '• Linux (Ubuntu 18.04+, Debian 10+)\n\n'
        '一套代码，多平台编译，确保在不同操作系统上获得一致的用户体验。'
    )
    
    doc.add_heading('离线运行能力', level=3)
    doc.add_paragraph(
        '系统具备完整的离线运行能力，核心设计如下：\n\n'
        '1. 本地数据存储：采用SQLite嵌入式数据库，所有业务数据存储在本地\n'
        '2. 内置后端服务：Python Flask服务与Electron应用打包为一体\n'
        '3. 独立运行：无需外部服务器，安装即可使用\n'
        '4. 数据安全：敏感数据不离开本地环境\n\n'
        '仅AI智能分析功能需要网络连接调用云端API，且发送的数据均经过脱敏处理。'
    )

def add_section_features(doc):
    """第三章：核心功能模块"""
    doc.add_heading('三、核心功能模块', level=1)
    
    doc.add_heading('3.1 数据管理模块', level=2)
    doc.add_paragraph(
        '系统提供完善的数据管理功能：\n\n'
        '• Excel批量导入：支持标准格式Excel文件一键导入\n'
        '• 数据校验：自动校验数据完整性和格式规范\n'
        '• 增删改查：完整的数据维护操作\n'
        '• 多维筛选：支持按时间、药品、科室等多维度筛选\n'
        '• 分页浏览：大数据量高效分页展示\n'
        '• 数据导出：支持导出为Excel格式'
    )
    
    doc.add_heading('3.2 可视化图表分析', level=2)
    doc.add_paragraph(
        '系统集成ECharts图表库，提供丰富的数据可视化能力：'
    )
    
    # 图表功能表格
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ['图表类型', '应用场景', '交互特性']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'DBEAFE')
    
    chart_data = [
        ('柱状图', '报告类型分布、药品TOP排名', '悬停提示、点击下钻'),
        ('饼图/环形图', '比例分析、类型占比', '扇区高亮、图例切换'),
        ('折线图', '月度趋势、同比环比', '数据缩放、区域选择'),
        ('堆叠图', '多维度对比分析', '系列切换、堆叠展开'),
        ('热力图', '时间-类别分布', '颜色映射、区域缩放'),
        ('组合图', '复杂多指标分析', '双Y轴、混合展示'),
    ]
    
    for i, row_data in enumerate(chart_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph(
        '图表特性：\n'
        '• 响应式布局：自适应不同屏幕尺寸\n'
        '• 动态数据：实时响应筛选条件变化\n'
        '• 导出功能：支持导出为PNG图片\n'
        '• 主题切换：支持多种配色方案'
    )
    
    doc.add_heading('3.3 AI智能助手', level=2)
    doc.add_paragraph(
        '系统集成AI智能助手，基于DeepSeek大语言模型，提供智能化分析能力：'
    )
    
    # AI功能表格
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ['功能模块', '功能描述', '输出形式']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'D1FAE5')
    
    ai_data = [
        ('数据洞察分析', '自动分析统计数据，发现规律和异常', 'Markdown报告'),
        ('趋势预测预警', '预测不良反应趋势，识别预警信号', '趋势图表+建议'),
        ('智能报告生成', '自动生成月度/季度分析报告', 'Word文档'),
        ('用药安全建议', '基于数据给出专业用药建议', '结构化建议'),
        ('智能图表生成', '自然语言描述生成统计图表', 'ECharts图表'),
    ]
    
    for i, row_data in enumerate(ai_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('3.4 工作量统计分析', level=2)
    doc.add_paragraph(
        '系统提供完整的工作量统计与绩效分析功能：\n\n'
        '• 工作量数据导入：支持Excel批量导入工作记录\n'
        '• 多维度统计：按人员、科室、时间段统计\n'
        '• 可视化报表：柱状图、饼图、趋势图等多种展示\n'
        '• 绩效考核：支持自定义考核指标\n'
        '• 报表导出：一键导出统计报表'
    )

def add_section_innovation(doc):
    """第四章：创新功能亮点"""
    doc.add_heading('四、创新功能亮点', level=1)
    
    doc.add_heading('4.1 AI驱动的智能分析', level=2)
    doc.add_paragraph(
        '系统创新性地将大语言模型（LLM）技术应用于药物警戒领域，实现了：\n\n'
        '自然语言交互：\n'
        '用户可通过自然语言描述分析需求，如"生成报告类型饼状图"、"分析药品TOP10"，'
        'AI自动理解意图并生成相应的统计图表或分析报告。\n\n'
        '智能数据解读：\n'
        'AI能够自动分析统计数据，识别异常值、发现趋势规律，并给出专业的药物警戒建议，'
        '大幅降低数据分析门槛。\n\n'
        '报告自动生成：\n'
        '支持一键生成结构化的分析报告，包含数据概览、趋势分析、风险预警、改进建议等章节，'
        '可直接导出为Word文档。'
    )
    
    doc.add_heading('4.2 数据隐私保护机制', level=2)
    doc.add_paragraph(
        '系统在设计之初即将数据隐私保护作为核心原则，创新性地实现了"隐私优先"的AI分析模式：'
    )
    
    # 隐私保护表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    headers = ['保护机制', '具体措施']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'FEE2E2')
    
    privacy_data = [
        ('敏感字段隔离', '报告编码、病历号、患者信息等敏感字段绝不发送至AI服务'),
        ('本地聚合计算', '所有统计计算在本地完成，仅发送聚合后的统计数值'),
        ('数据脱敏验证', '发送前自动验证数据包，确保不含敏感信息'),
        ('用户透明告知', '界面明确提示"AI不访问患者隐私信息"'),
    ]
    
    for i, row_data in enumerate(privacy_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph(
        '敏感字段黑名单：\n'
        '• report_code（报告编码）\n'
        '• medical_record_no（病历号）\n'
        '• reporter_signature（报告人签名）\n'
        '• 任何可识别患者身份的信息\n\n'
        '发送给AI的安全数据仅包含：统计数量、比例、趋势等聚合指标。'
    )
    
    doc.add_heading('4.3 交互式可视化图表', level=2)
    doc.add_paragraph(
        '系统采用ECharts图表库，实现了高度交互的数据可视化体验：\n\n'
        '• 动态筛选：图表与筛选条件联动，实时更新\n'
        '• 缩放平移：支持数据区域的缩放和平移操作\n'
        '• 数据下钻：点击图表元素可查看详细数据\n'
        '• 图例交互：点击图例切换数据系列显示\n'
        '• 工具箱：内置保存图片、数据视图等工具\n'
        '• 响应式设计：自适应不同窗口尺寸'
    )

def add_section_security(doc):
    """第五章：数据安全与隐私保护"""
    doc.add_heading('五、数据安全与隐私保护', level=1)
    
    doc.add_heading('5.1 安全设计原则', level=2)
    doc.add_paragraph(
        '系统遵循"隐私优先、最小必要、本地优先"的安全设计原则：\n\n'
        '1. 隐私优先：任何情况下都不将敏感数据发送至外部服务\n'
        '2. 最小必要：AI分析仅使用完成任务所必需的最小数据集\n'
        '3. 本地优先：核心业务数据存储于本地，离线可用\n'
        '4. 透明可控：用户清楚知道哪些数据被使用'
    )
    
    doc.add_heading('5.2 数据流安全', level=2)
    
    flow_para = doc.add_paragraph()
    flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = flow_para.add_run(
        '┌──────────────┐     ┌──────────────┐     ┌──────────────┐\n'
        '│   原始数据库    │ ──→ │   聚合统计计算   │ ──→ │   安全数据包    │\n'
        '│  (含敏感字段)   │     │   (本地处理)    │     │   (仅统计值)    │\n'
        '└──────────────┘     └──────────────┘     └──────────────┘\n'
        '                                                    │\n'
        '                                                    ↓\n'
        '                                          ┌──────────────┐\n'
        '                                          │  DeepSeek API │\n'
        '                                          │   (云端AI)    │\n'
        '                                          └──────────────┘'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('5.3 认证与授权', level=2)
    doc.add_paragraph(
        '• 用户认证：基于JWT Token的身份认证机制\n'
        '• 接口保护：所有API接口均需携带有效Token\n'
        '• 权限控制：支持角色权限管理（可扩展）\n'
        '• 会话管理：支持自动登出和会话超时'
    )

def add_section_deployment(doc):
    """第六章：系统部署与运行"""
    doc.add_heading('六、系统部署与运行', level=1)
    
    doc.add_heading('6.1 部署方式', level=2)
    doc.add_paragraph(
        '系统支持两种部署方式：\n\n'
        '独立安装包部署（推荐）：\n'
        '• 下载对应操作系统的安装包\n'
        '• Windows: .exe安装程序\n'
        '• macOS: .dmg安装镜像\n'
        '• Linux: .deb/.rpm安装包\n'
        '• 双击安装，开箱即用\n\n'
        '开发环境部署：\n'
        '• 克隆代码仓库\n'
        '• 安装Node.js和Python环境\n'
        '• 安装依赖：npm install / pip install\n'
        '• 启动服务：npm start'
    )
    
    doc.add_heading('6.2 系统要求', level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['项目', '最低要求', '推荐配置']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'E0E7FF')
    
    req_data = [
        ('操作系统', 'Windows 10 / macOS 10.15 / Ubuntu 18.04', 'Windows 11 / macOS 14 / Ubuntu 22.04'),
        ('内存', '4GB RAM', '8GB+ RAM'),
        ('存储空间', '500MB可用空间', '1GB+ 可用空间'),
        ('网络', '可选（AI功能需要）', '稳定网络连接'),
    ]
    
    for i, row_data in enumerate(req_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i + 1].cells[j].text = cell_text

def add_section_summary(doc):
    """第七章：总结与展望"""
    doc.add_heading('七、总结与展望', level=1)
    
    doc.add_heading('7.1 系统特色总结', level=2)
    doc.add_paragraph(
        '医院药学药物警戒与运营绩效可视化分析系统具有以下核心特色：\n\n'
        '✓ 跨平台支持：一套代码，Windows/macOS/Linux全覆盖\n'
        '✓ 离线运行：核心功能无需网络，数据本地存储\n'
        '✓ AI智能化：大语言模型赋能，智能分析与报告生成\n'
        '✓ 隐私保护：敏感数据不出本地，AI仅接收统计数据\n'
        '✓ 可视化丰富：ECharts驱动，交互式数据图表\n'
        '✓ 易于使用：现代化UI设计，操作简单直观'
    )
    
    doc.add_heading('7.2 未来展望', level=2)
    doc.add_paragraph(
        '系统后续可扩展的方向：\n\n'
        '• 多院区数据整合：支持集团医院多院区数据汇总\n'
        '• 智能预警升级：基于机器学习的风险预测模型\n'
        '• 移动端支持：开发配套的移动App\n'
        '• 监管对接：与国家药品不良反应监测系统对接\n'
        '• 知识图谱：构建药物警戒领域知识图谱'
    )
    
    doc.add_paragraph()
    
    # 结束语
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run('— 文档结束 —')
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

def main():
    print('正在生成系统技术文档...')
    
    doc = create_document()
    
    # 添加各章节
    add_cover_page(doc)
    add_toc(doc)
    add_section_overview(doc)
    add_section_architecture(doc)
    add_section_features(doc)
    add_section_innovation(doc)
    add_section_security(doc)
    add_section_deployment(doc)
    add_section_summary(doc)
    
    # 保存文档
    output_path = '医院药学药物警戒与运营绩效可视化分析系统_技术文档.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')

if __name__ == '__main__':
    main()
